"""
Document Parser Tools
Extracts text from various document formats (PDF, DOCX, TXT)
"""
import io
from typing import Dict, Any, Optional
from pathlib import Path


class DocumentParser:
    """Unified document parser for multiple formats"""
    
    @staticmethod
    def detect_format(file_path: str = None, file_bytes: bytes = None, filename: str = None) -> str:
        """
        Detect document format from file extension or content
        
        Args:
            file_path: Path to file
            file_bytes: File content as bytes
            filename: Original filename
            
        Returns:
            Format: 'pdf', 'docx', 'txt', or 'unknown'
        """
        if file_path:
            ext = Path(file_path).suffix.lower()
        elif filename:
            ext = Path(filename).suffix.lower()
        else:
            # Try to detect from bytes
            if file_bytes:
                if file_bytes.startswith(b'%PDF'):
                    return 'pdf'
                elif file_bytes.startswith(b'PK'):  # ZIP format (DOCX)
                    return 'docx'
            return 'unknown'
        
        format_map = {
            '.pdf': 'pdf',
            '.docx': 'docx',
            '.doc': 'doc',
            '.txt': 'txt',
            '.md': 'txt',
        }
        
        return format_map.get(ext, 'unknown')
    
    @staticmethod
    async def parse_pdf(file_bytes: bytes) -> Dict[str, Any]:
        """
        Extract text from PDF file
        
        Args:
            file_bytes: PDF file content
            
        Returns:
            Dictionary with text and metadata
        """
        try:
            import pymupdf  # PyMuPDF
            
            # Open PDF from bytes
            doc = pymupdf.open(stream=file_bytes, filetype="pdf")
            
            text_content = []
            metadata = {
                "pages": len(doc),
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", ""),
            }
            
            # Extract text from each page
            for page_num in range(len(doc)):
                page = doc[page_num]
                text_content.append(page.get_text())
            
            doc.close()
            
            return {
                "text": "\n\n".join(text_content),
                "metadata": metadata,
                "format": "pdf"
            }
            
        except Exception as e:
            return {
                "text": "",
                "metadata": {},
                "format": "pdf",
                "error": str(e)
            }
    
    @staticmethod
    async def parse_docx(file_bytes: bytes) -> Dict[str, Any]:
        """
        Extract text from DOCX file
        
        Args:
            file_bytes: DOCX file content
            
        Returns:
            Dictionary with text and metadata
        """
        try:
            from docx import Document
            
            # Open DOCX from bytes
            doc = Document(io.BytesIO(file_bytes))
            
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            metadata = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
            }
            
            # Try to get core properties
            try:
                core_props = doc.core_properties
                metadata["author"] = core_props.author or ""
                metadata["title"] = core_props.title or ""
            except:
                pass
            
            return {
                "text": "\n".join(text_content),
                "metadata": metadata,
                "format": "docx"
            }
            
        except Exception as e:
            return {
                "text": "",
                "metadata": {},
                "format": "docx",
                "error": str(e)
            }
    
    @staticmethod
    async def parse_text(file_bytes: bytes) -> Dict[str, Any]:
        """
        Extract text from plain text file
        
        Args:
            file_bytes: Text file content
            
        Returns:
            Dictionary with text and metadata
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            text = None
            
            for encoding in encodings:
                try:
                    text = file_bytes.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                text = file_bytes.decode('utf-8', errors='ignore')
            
            return {
                "text": text,
                "metadata": {
                    "size_bytes": len(file_bytes),
                    "lines": len(text.split('\n'))
                },
                "format": "txt"
            }
            
        except Exception as e:
            return {
                "text": "",
                "metadata": {},
                "format": "txt",
                "error": str(e)
            }
    
    @classmethod
    async def parse_document(
        cls,
        file_bytes: bytes,
        filename: Optional[str] = None,
        force_format: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Universal document parser
        
        Args:
            file_bytes: Document content as bytes
            filename: Original filename (for format detection)
            force_format: Force specific format ('pdf', 'docx', 'txt')
            
        Returns:
            Parsed document with text and metadata
        """
        # Detect format
        if force_format:
            doc_format = force_format
        else:
            doc_format = cls.detect_format(file_bytes=file_bytes, filename=filename)
        
        # Parse based on format
        if doc_format == 'pdf':
            return await cls.parse_pdf(file_bytes)
        elif doc_format == 'docx' or doc_format == 'doc':
            return await cls.parse_docx(file_bytes)
        elif doc_format == 'txt':
            return await cls.parse_text(file_bytes)
        else:
            return {
                "text": "",
                "metadata": {},
                "format": "unknown",
                "error": f"Unsupported format: {doc_format}"
            }


# Convenience functions
async def parse_cv_document(file_bytes: bytes, filename: str = None) -> str:
    """
    Parse CV document and return text content
    
    Args:
        file_bytes: Document bytes
        filename: Original filename
        
    Returns:
        Extracted text
    """
    result = await DocumentParser.parse_document(file_bytes, filename)
    
    if result.get("error"):
        raise ValueError(f"Failed to parse document: {result['error']}")
    
    return result["text"]
