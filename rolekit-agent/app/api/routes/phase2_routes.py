"""
Phase 2 API Routes - Node-based CV Processing Endpoints
Each endpoint corresponds to a specific processing node in the CV enhancement pipeline.
"""
from fastapi import APIRouter, UploadFile, File, Form, HTTPException, BackgroundTasks
from fastapi.responses import Response, FileResponse, JSONResponse
from typing import Optional, Literal
from pydantic import BaseModel, Field
import tempfile
import os
from pathlib import Path
import uuid
from datetime import datetime

from app.models.cv_models import CVData
from app.services.cv import (
    CVSchemaExtractor,
    CVSchemaValidator,
    ProfileEnhancer,
    ImpactQuantifier,
    CVBuilder
)
from app.services.parser.document_parser import DocumentParser
from app.services.pdf_generator import PDFGenerator, PDFGenerationError
from app.core.dependencies import get_llm

# Create router with /api prefix
router = APIRouter(prefix="/api", tags=["CV Processing Pipeline"])

# Directory for temporary exports
EXPORT_DIR = Path("exports")
EXPORT_DIR.mkdir(exist_ok=True)


# ============================================================================
# REQUEST/RESPONSE MODELS
# ============================================================================

class ExtractRequest(BaseModel):
    """Request model for /api/extract endpoint"""
    text: Optional[str] = Field(None, description="Raw CV text")
    target_role: Optional[str] = Field(None, description="Target job role for context")
    extract_skills: bool = Field(True, description="Extract skills section")
    extract_education: bool = Field(True, description="Extract education section")


class ExtractResponse(BaseModel):
    """Response model for /api/extract endpoint"""
    success: bool
    cv_data: dict
    confidence_score: Optional[float] = None
    warnings: Optional[list[str]] = None
    message: str


class EnhanceRequest(BaseModel):
    """Request model for /api/enhance endpoint"""
    cv_data: CVData
    target_role: Optional[str] = Field(None, description="Target role for optimization")
    enhancement_focus: list[str] = Field(
        default=["clarity", "impact", "keywords", "tone"],
        description="Areas to focus enhancement on"
    )
    enhance_summary: bool = Field(True, description="Enhance professional summary")
    enhance_experience: bool = Field(True, description="Enhance work experience")
    enhance_project: bool = Field(False, description="Enhance projects")
    quantify_achievements: bool = Field(True, description="Add metrics to achievements")


class EnhanceResponse(BaseModel):
    """Response model for /api/enhance endpoint"""
    success: bool
    enhanced_cv: dict
    improvements: dict = Field(description="Summary of improvements made")
    before_after_comparison: Optional[dict] = None
    message: str


class BuildRequest(BaseModel):
    """Request model for /api/build endpoint"""
    cv_data: CVData
    format: Literal["html", "markdown", "json"] = Field(
        default="html",
        description="Output format"
    )
    style: str = Field(
        default="modern",
        description="Template style (modern, classic, minimal)"
    )
    include_photo: bool = Field(False, description="Include profile photo")
    color_scheme: Optional[str] = Field("blue", description="Color scheme for HTML")


class BuildResponse(BaseModel):
    """Response model for /api/build endpoint"""
    success: bool
    content: str = Field(description="Generated CV content")
    format: str
    preview_url: Optional[str] = None
    message: str


class ExportRequest(BaseModel):
    """Request model for /api/export endpoint"""
    cv_data: CVData
    format: Literal["pdf", "docx"] = Field(default="pdf", description="Export format")
    style: str = Field(default="modern", description="Template style")


class ExportResponse(BaseModel):
    """Response model for /api/export endpoint"""
    success: bool
    download_url: str
    file_id: str
    expires_at: str
    message: str


class FeedbackRequest(BaseModel):
    """Request model for /api/feedback endpoint"""
    original_cv: CVData
    improved_cv: CVData
    criteria: list[str] = Field(
        default=["clarity", "impact", "professionalism", "ats_compatibility"],
        description="Evaluation criteria"
    )


class FeedbackResponse(BaseModel):
    """Response model for /api/feedback endpoint"""
    success: bool
    overall_score: float = Field(description="Overall improvement score (0-100)")
    detailed_feedback: dict = Field(description="Feedback by section")
    strengths: list[str] = Field(description="Improvements made well")
    weaknesses: list[str] = Field(description="Areas needing more work")
    recommendations: list[str] = Field(description="Suggestions for further improvement")
    message: str


class SuggestSkillsRequest(BaseModel):
    """Request model for /api/suggest-skills endpoint"""
    cv_data: dict = Field(description="CV data with experience, projects, and technologies")


class SuggestSkillsResponse(BaseModel):
    """Response model for /api/suggest-skills endpoint"""
    success: bool
    suggested_skills: list[str] = Field(description="List of suggested skills")
    count: int = Field(description="Number of suggested skills")
    message: str


class ValidateSkillRequest(BaseModel):
    """Request model for /api/validate-skill endpoint"""
    skill: str = Field(description="Raw skill name to validate")
    existing_skills: list[str] = Field(default=[], description="Existing skills to avoid duplicates")


class ValidateSkillResponse(BaseModel):
    """Response model for /api/validate-skill endpoint"""
    success: bool
    validated_skill: str = Field(description="Corrected and validated skill name")
    correction_applied: bool = Field(description="Whether the skill was corrected")
    message: str


# ============================================================================
# ENDPOINT 1: /api/extract - Parse and Extract CV Data
# ============================================================================

@router.post("/extract", response_model=ExtractResponse)
async def extract_cv_data(request: ExtractRequest):
    """
    **Node: Input Parser + Schema Extractor**
    
    Accepts file/text → parses → returns JSON structure
    
    This endpoint:
    1. Accepts raw CV text or file
    2. Uses LLM to extract structured data
    3. Validates and cleans the data
    4. Returns JSON with confidence scores
    
    **Example:**
    ```json
    {
        "text": "John Doe\\nSoftware Engineer\\n...",
        "target_role": "Senior Developer"
    }
    ```
    """
    try:
        if not request.text:
            raise HTTPException(status_code=400, detail="Text is required")
        
        # Initialize extractor
        extractor = CVSchemaExtractor()
        validator = CVSchemaValidator()
        
        # Extract with or without context
        if request.target_role:
            cv_data = await extractor.extract_with_context(
                request.text,
                target_role=request.target_role
            )
        else:
            cv_data = await extractor.extract(request.text)
        
        # Validate and get quality score
        cv_data = validator.validate_and_clean(cv_data)
        quality_result = await validator.validate_with_llm(cv_data)
        
        # Extract quality score
        quality_score = quality_result.get("quality_score", 0.5) / 100.0  # Convert 0-100 to 0-1
        
        # Collect warnings
        warnings = []
        if quality_score < 0.7:
            warnings.append("Low confidence in extracted data. Please review carefully.")
        if not cv_data.contact or not cv_data.contact.email:
            warnings.append("Email address not found")
        if not cv_data.experience or len(cv_data.experience) == 0:
            warnings.append("No work experience found")
        
        return ExtractResponse(
            success=True,
            cv_data=cv_data.model_dump(),
            confidence_score=quality_score,
            warnings=warnings if warnings else None,
            message=f"CV data extracted successfully with {quality_score:.0%} confidence"
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Extraction failed: {str(e)}")


@router.post("/extract/upload", response_model=ExtractResponse)
async def extract_cv_from_file(
    file: UploadFile = File(...),
    target_role: Optional[str] = Form(None)
):
    """
    **Node: Input Parser + Schema Extractor**
    
    Upload file version of /api/extract
    
    Accepts: PDF, DOCX, TXT files
    """
    try:
        # Read file content
        content = await file.read()
        
        # Parse document
        parser = DocumentParser()
        result = await parser.parse_document(content, file.filename)
        text = result.get("text", "")
        
        if not text or len(text.strip()) < 50:
            raise HTTPException(
                status_code=400,
                detail="Could not extract meaningful text from file"
            )
        
        # Use the text extraction endpoint
        request = ExtractRequest(text=text, target_role=target_role)
        return await extract_cv_data(request)
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"File processing failed: {str(e)}"
        )


# ============================================================================
# ENDPOINT 2: /api/enhance - Enhance CV Content
# ============================================================================

@router.post("/enhance", response_model=EnhanceResponse)
async def enhance_cv_content(request: EnhanceRequest):
    """
    **Node: Profile Enhancer + Impact Quantifier**
    
    Takes CV JSON → rewrites text → returns improved version
    
    This endpoint:
    1. Enhances professional summary
    2. Improves work experience descriptions
    3. Quantifies achievements with metrics
    4. Optimizes language for impact
    5. Maintains professional tone
    
    **Example:**
    ```json
    {
        "cv_data": {...},
        "target_role": "Senior Developer",
        "enhancement_focus": ["clarity", "impact", "keywords"]
    }
    ```
    """
    try:
        enhancer = ProfileEnhancer()
        quantifier = ImpactQuantifier()
        
        # Store original for comparison
        original_cv = request.cv_data.model_copy(deep=True)
        enhanced_cv = request.cv_data.model_copy(deep=True)
        
        improvements = {
            "summary_enhanced": False,
            "experiences_enhanced": 0,
            "metrics_added": 0,
            "keywords_optimized": False
        }
        
        # 1. Enhance Professional Summary
        if request.enhance_summary and enhanced_cv.summary:
            original_summary = enhanced_cv.summary
            enhanced_cv.summary = await enhancer.enhance_summary(
                enhanced_cv.summary,
                role=request.target_role
            )
            if enhanced_cv.summary != original_summary:
                improvements["summary_enhanced"] = True
        
        # 2. Enhance Work Experience
        if request.enhance_experience and enhanced_cv.experience:
            for i, exp in enumerate(enhanced_cv.experience):
                # Enhance description
                enhanced_exp = await enhancer.enhance_experience_description(
                    exp,
                    focus_on=request.enhancement_focus
                )
                enhanced_cv.experience[i] = enhanced_exp
                improvements["experiences_enhanced"] += 1
        
        # 2b. Enhance Education Degrees
        if enhanced_cv.education:
            for i, edu in enumerate(enhanced_cv.education):
                # Convert to dict for enhancement
                edu_dict = {
                    "institution": edu.institution or "",
                    "degree": edu.degree or "",
                    "field_of_study": edu.field_of_study or "",
                    "school": edu.institution or ""
                }
                # Enhance the degree
                enhanced_edu_dict = await enhancer.enhance_education_degree(edu_dict)
                # Update the degree field
                if enhanced_edu_dict.get("degree"):
                    enhanced_cv.education[i].degree = enhanced_edu_dict["degree"]
        
        # 2c. Enhance Projects (description and technologies)
        if request.enhance_project and enhanced_cv.projects:
            for i, proj in enumerate(enhanced_cv.projects):
                # Convert to dict for enhancement
                proj_dict = {
                    "name": proj.name or "",
                    "description": proj.description or "",
                    "technologies": proj.technologies or [],
                    "url": proj.url or "",
                    "repository": proj.repository or ""
                }
                # Enhance the project
                enhanced_proj_dict = await enhancer.enhance_project(proj_dict)
                # Update the project fields
                if enhanced_proj_dict.get("description"):
                    enhanced_cv.projects[i].description = enhanced_proj_dict["description"]
                if enhanced_proj_dict.get("technologies"):
                    enhanced_cv.projects[i].technologies = enhanced_proj_dict["technologies"]
        
        # 3. Optimize keywords if requested
        if "keywords" in request.enhancement_focus and request.target_role:
            improvements["keywords_optimized"] = True
        
        # 4. Create before/after comparison
        before_after = {
            "summary": {
                "before": original_cv.summary,
                "after": enhanced_cv.summary
            },
            "experience_count": len(enhanced_cv.experience or []),
            "improvements_made": improvements["experiences_enhanced"]
        }
        
        return EnhanceResponse(
            success=True,
            enhanced_cv=enhanced_cv.model_dump(),
            improvements=improvements,
            before_after_comparison=before_after,
            message=f"CV enhanced successfully. Improved {improvements['experiences_enhanced']} experiences, added {improvements['metrics_added']} metrics"
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Enhancement failed: {str(e)}")


# ============================================================================
# ENDPOINT 3: /api/build - Build Styled CV
# ============================================================================

@router.post("/build", response_model=BuildResponse)
async def build_styled_cv(request: BuildRequest):
    """
    **Node: CV Builder**
    
    Converts JSON → styled Markdown/HTML CV
    
    This endpoint:
    1. Takes structured CV data
    2. Applies chosen template and style
    3. Generates formatted output (HTML, Markdown, or JSON)
    4. Returns ready-to-use content
    
    **Example:**
    ```json
    {
        "cv_data": {...},
        "format": "html",
        "style": "modern",
        "color_scheme": "blue"
    }
    ```
    """
    try:
        builder = CVBuilder()
        
        # Generate content based on format
        if request.format == "html":
            content = builder.to_html(
                request.cv_data,
                style=request.style
            )
            
        elif request.format == "markdown":
            content = builder.to_markdown(request.cv_data)
            
        elif request.format == "json":
            content = builder.to_json(request.cv_data)
            
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported format: {request.format}"
            )
        
        # For HTML, we could save and provide preview URL
        preview_url = None
        if request.format == "html":
            # Save to temp location for preview
            file_id = str(uuid.uuid4())
            preview_path = EXPORT_DIR / f"preview_{file_id}.html"
            preview_path.write_text(content, encoding="utf-8")
            preview_url = f"/api/preview/{file_id}"
        
        return BuildResponse(
            success=True,
            content=content,
            format=request.format,
            preview_url=preview_url,
            message=f"CV built successfully in {request.format.upper()} format"
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Build failed: {str(e)}")


# ============================================================================
# ENDPOINT 4: /api/export - Export to PDF/DOCX
# ============================================================================

@router.post("/export", response_model=ExportResponse)
async def export_cv_file(request: ExportRequest, background_tasks: BackgroundTasks):
    """
    **Node: PDF/DOCX Generator**
    
    Converts HTML → PDF/DOCX → download link
    
    This endpoint:
    1. Takes CV data
    2. Generates HTML
    3. Converts to PDF or DOCX
    4. Provides download link
    5. Auto-cleanup after 1 hour
    
    **Example:**
    ```json
    {
        "cv_data": {...},
        "format": "pdf",
        "style": "modern"
    }
    ```
    """
    try:
        builder = CVBuilder()
        
        # Generate unique file ID
        file_id = str(uuid.uuid4())
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Build HTML first
        html_content = builder.to_html(request.cv_data, style=request.style)
        
        if request.format == "pdf":
            # Use the PDFGenerator service
            filename = f"cv_{timestamp}_{file_id}.pdf"
            filepath = EXPORT_DIR / filename
            
            try:
                pdf_generator = PDFGenerator()
                pdf_generator.generate_pdf(html_content, filepath)
            except PDFGenerationError as e:
                raise HTTPException(
                    status_code=500,
                    detail=f"PDF generation failed: {str(e)}. "
                           "Ensure wkhtmltopdf is installed or weasyprint is available."
                )
            
        elif request.format == "docx":
            # Convert to DOCX
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            filename = f"cv_{timestamp}_{file_id}.docx"
            filepath = EXPORT_DIR / filename
            
            doc = Document()
            
            # Add contact info
            if request.cv_data.contact_info:
                contact = request.cv_data.contact_info
                heading = doc.add_heading(contact.full_name or "Professional CV", 0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                if contact.email or contact.phone:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if contact.email:
                        p.add_run(f"Email: {contact.email}  ")
                    if contact.phone:
                        p.add_run(f"Phone: {contact.phone}")
            
            # Add summary
            if request.cv_data.summary:
                doc.add_heading('Professional Summary', 1)
                doc.add_paragraph(request.cv_data.summary)
            
            # Add work experience
            if request.cv_data.experience:
                doc.add_heading('Work Experience', 1)
                for exp in request.cv_data.experience:
                    doc.add_heading(f"{exp.position} at {exp.company}", 2)
                    p = doc.add_paragraph(f"{exp.start_date} - {exp.end_date or 'Present'}")
                    if exp.description:
                        doc.add_paragraph(exp.description)
                    if exp.achievements:
                        for achievement in exp.achievements:
                            doc.add_paragraph(achievement, style='List Bullet')
            
            # Add education
            if request.cv_data.education:
                doc.add_heading('Education', 1)
                for edu in request.cv_data.education:
                    doc.add_heading(f"{edu.degree} - {edu.institution}", 2)
                    doc.add_paragraph(f"Graduated: {edu.graduation_date or 'In Progress'}")
            
            # Add skills
            if request.cv_data.skills:
                doc.add_heading('Skills', 1)
                for skill in request.cv_data.skills:
                    doc.add_paragraph(f"{skill.name} - {skill.level or 'Proficient'}", style='List Bullet')
            
            # Save document
            doc.save(str(filepath))
        
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported export format: {request.format}"
            )
        
        # Schedule cleanup after 1 hour
        def cleanup_file():
            import time
            time.sleep(3600)  # 1 hour
            if filepath.exists():
                filepath.unlink()
        
        background_tasks.add_task(cleanup_file)
        
        # Generate expiry time
        from datetime import timedelta
        expires_at = (datetime.now() + timedelta(hours=1)).isoformat()
        
        return ExportResponse(
            success=True,
            download_url=f"/api/download/{file_id}",
            file_id=file_id,
            expires_at=expires_at,
            message=f"CV exported successfully to {request.format.upper()}"
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")


@router.get("/download/{file_id}")
async def download_exported_file(file_id: str):
    """Download exported CV file"""
    try:
        # Find file with this ID
        files = list(EXPORT_DIR.glob(f"*{file_id}*"))
        
        if not files:
            raise HTTPException(status_code=404, detail="File not found or expired")
        
        filepath = files[0]
        
        # Determine media type
        if filepath.suffix == ".pdf":
            media_type = "application/pdf"
        elif filepath.suffix == ".docx":
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            media_type = "application/octet-stream"
        
        return FileResponse(
            path=str(filepath),
            media_type=media_type,
            filename=filepath.name
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Download failed: {str(e)}")


@router.get("/preview/{file_id}")
async def preview_cv(file_id: str):
    """Preview HTML CV in browser"""
    try:
        filepath = EXPORT_DIR / f"preview_{file_id}.html"
        
        if not filepath.exists():
            raise HTTPException(status_code=404, detail="Preview not found or expired")
        
        return FileResponse(
            path=str(filepath),
            media_type="text/html"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Preview failed: {str(e)}")


# ============================================================================
# ENDPOINT 5: /api/feedback - Compare CV Versions
# ============================================================================

@router.post("/feedback", response_model=FeedbackResponse)
async def compare_cv_versions(request: FeedbackRequest):
    """
    **Node: Feedback & Iteration**
    
    Compare two CV versions using LLM
    
    This endpoint:
    1. Compares original vs improved CV
    2. Evaluates improvement across criteria
    3. Provides detailed feedback by section
    4. Suggests further improvements
    5. Generates overall quality score
    
    **Example:**
    ```json
    {
        "original_cv": {...},
        "improved_cv": {...},
        "criteria": ["clarity", "impact", "professionalism"]
    }
    ```
    """
    try:
        from langchain_core.prompts import ChatPromptTemplate
        from langchain_core.output_parsers import JsonOutputParser
        
        llm = get_llm()
        
        # Build comparison prompt
        prompt = ChatPromptTemplate.from_messages([
            ("system", """You are an expert CV reviewer. Compare two versions of a CV and provide detailed feedback.

Evaluate the improvements based on these criteria:
{criteria}

Provide your analysis in JSON format with:
- overall_score: 0-100 score for improvement
- detailed_feedback: object with feedback for each section (summary, experience, education, skills)
- strengths: array of things done well
- weaknesses: array of areas needing improvement
- recommendations: array of specific suggestions

Be constructive, specific, and actionable."""),
            ("user", """
**ORIGINAL CV:**
Summary: {original_summary}
Work Experience: {original_experience}
Skills: {original_skills}

**IMPROVED CV:**
Summary: {improved_summary}
Work Experience: {improved_experience}
Skills: {improved_skills}

Provide detailed comparison and feedback.""")
        ])
        
        # Prepare data for comparison
        original_summary = request.original_cv.summary or "Not provided"
        improved_summary = request.improved_cv.summary or "Not provided"
        
        original_experience = "\n".join([
            f"- {exp.position} at {exp.company}: {exp.description or 'No description'}"
            for exp in (request.original_cv.experience or [])
        ]) or "Not provided"
        
        improved_experience = "\n".join([
            f"- {exp.position} at {exp.company}: {exp.description or 'No description'}"
            for exp in (request.improved_cv.experience or [])
        ]) or "Not provided"
        
        original_skills = ", ".join([
            skill.name for skill in (request.original_cv.skills or [])
        ]) or "Not provided"
        
        improved_skills = ", ".join([
            skill.name for skill in (request.improved_cv.skills or [])
        ]) or "Not provided"
        
        # Run LLM analysis
        chain = prompt | llm | JsonOutputParser()
        
        result = await chain.ainvoke({
            "criteria": ", ".join(request.criteria),
            "original_summary": original_summary,
            "improved_summary": improved_summary,
            "original_experience": original_experience,
            "improved_experience": improved_experience,
            "original_skills": original_skills,
            "improved_skills": improved_skills
        })
        
        return FeedbackResponse(
            success=True,
            overall_score=float(result.get("overall_score", 75)),
            detailed_feedback=result.get("detailed_feedback", {}),
            strengths=result.get("strengths", []),
            weaknesses=result.get("weaknesses", []),
            recommendations=result.get("recommendations", []),
            message="CV comparison completed successfully"
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Feedback generation failed: {str(e)}")


@router.post("/suggest-skills", response_model=SuggestSkillsResponse)
async def suggest_skills(request: SuggestSkillsRequest):
    """
    Suggest relevant skills based on experience, projects, and technologies.
    
    Uses LLM to analyze job titles, project descriptions, and technologies
    to suggest skills that would strengthen the CV.
    
    Args:
        request: CV data containing experience and project information
        
    Returns:
        SuggestSkillsResponse with list of suggested skills
    """
    try:
        # Initialize profile enhancer
        enhancer = ProfileEnhancer()
        
        # Get skill suggestions
        suggested_skills = await enhancer.suggest_skills(request.cv_data)
        
        if suggested_skills:
            return SuggestSkillsResponse(
                success=True,
                suggested_skills=suggested_skills,
                count=len(suggested_skills),
                message=f"Successfully generated {len(suggested_skills)} skill suggestions based on your experience"
            )
        else:
            return SuggestSkillsResponse(
                success=True,
                suggested_skills=[],
                count=0,
                message="No new skills to suggest. Your CV already includes relevant skills!"
            )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Skill suggestion failed: {str(e)}")


# ============================================================================
# VALIDATE SKILL
# ============================================================================

@router.post("/validate-skill")
async def validate_skill(request: ValidateSkillRequest):
    """
    Validate if a given skill is a standard/recognized skill.
    
    Parameters:
    - skill (str): The skill to validate
    
    Returns validated skill status and recommendations
    """
    try:
        skill = request.skill.strip()
        if not skill:
            raise ValueError("Skill parameter is required")
        
        # List of recognized/standard skills
        standard_skills = {
            "python", "javascript", "java", "c++", "c#", "php", "ruby", "go", "rust", 
            "swift", "kotlin", "scala", "typescript", "groovy", "perl", "r", "matlab",
            "html", "css", "xml", "json", "sql", "nosql", "mongodb", "postgresql", 
            "mysql", "redis", "cassandra", "elasticsearch", "neo4j", "firebase",
            "react", "angular", "vue", "vue.js", "svelte", "next.js", "nuxt", "ember",
            "node.js", "express", "django", "flask", "fastapi", "spring", "spring boot",
            "asp.net", "asp.net core", ".net", "laravel", "symfony", "rails", "sinatra",
            "gin", "echo", "beego", "fasthttp", "actix-web", "iron", "warp",
            "docker", "kubernetes", "jenkins", "gitlab ci", "github actions", "circleci",
            "travis ci", "aws", "azure", "gcp", "google cloud", "heroku", "vercel",
            "netlify", "digitalocean", "linode", "terraform", "ansible", "puppet", "chef",
            "prometheus", "grafana", "splunk", "datadog", "new relic", "elastic stack",
            "git", "svn", "mercurial", "perforce", "jira", "confluence", "trello",
            "asana", "monday.com", "slack", "teams", "discord", "mattermost",
            "agile", "scrum", "kanban", "waterfall", "lean", "six sigma", "tdd", "bdd",
            "rest api", "graphql", "grpc", "soap", "websockets", "mqtt", "amqp",
            "machine learning", "deep learning", "nlp", "computer vision", "pytorch",
            "tensorflow", "keras", "scikit-learn", "pandas", "numpy", "scipy", "matplotlib",
            "seaborn", "plotly", "dash", "jupyter", "anaconda", "conda", "pip",
            "git", "github", "gitlab", "bitbucket", "gitea", "gitops",
            "linux", "unix", "windows", "macos", "ios", "android", "windows phone",
            "bash", "zsh", "powershell", "cmd", "shell scripting", "awk", "sed", "grep",
            "vim", "emacs", "vs code", "visual studio", "intellij idea", "pycharm",
            "sublime text", "atom", "eclipse", "netbeans", "xcode", "android studio",
            "figma", "sketch", "adobe xd", "photoshop", "illustrator", "lightroom",
            "blender", "unity", "unreal engine", "godot", "twine", "construct",
            "agile", "scrum", "kanban", "waterfall", "devops", "sre", "mlops",
            "security", "cryptography", "penetration testing", "ethical hacking",
            "owasp", "gdpr", "ccpa", "hipaa", "pci-dss", "iso 27001",
            "microservices", "monolith", "serverless", "edge computing", "5g",
            "blockchain", "ethereum", "smart contracts", "web3", "defi", "nft",
            "excel", "power bi", "tableau", "looker", "qlik", "microstrategy",
            "sap", "oracle", "salesforce", "sap", "erp", "crm", "hrm",
            "scrum master", "product owner", "architect", "devops engineer", "data scientist",
            "data engineer", "analyst", "tester", "qa", "ux", "ui", "designer",
            "technical writing", "documentation", "api documentation", "swagger", "openapi",
            "communication", "teamwork", "leadership", "problem-solving", "creativity",
            "critical thinking", "time management", "adaptability", "learning", "mentoring"
        }
        
        skill_lower = skill.lower()
        is_standard = skill_lower in standard_skills
        
        return {
            "success": True,
            "skill": skill,
            "is_standard": is_standard,
            "status": "Valid standard skill" if is_standard else "Custom skill (not standard but acceptable)",
            "message": f"Skill '{skill}' is recognized as a {'standard' if is_standard else 'custom'} skill"
        }
        
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Skill validation failed: {str(e)}")


# ============================================================================
# FUTURE GOALS GENERATION
# ============================================================================

@router.post("/generate-future-goals")
async def generate_future_goals(request: dict):
    """
    Generate future career goals based on professional summary and experience.
    
    Analyzes the user's background and generates personalized career aspirations.
    """
    try:
        text = request.get("text", "")
        
        if not text or not text.strip():
            raise HTTPException(status_code=400, detail="No content provided to generate goals")
        
        enhancer = ProfileEnhancer()
        
        # Generate future goals using the LLM's enhance_summary capability
        # This method will intelligently generate career goals based on the input
        future_goals = await enhancer.enhance_summary(text)
        
        return {
            "success": True,
            "enhanced_text": future_goals,
            "message": "Future goals generated successfully"
        }
        
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to generate future goals: {str(e)}")


# ============================================================================
# HEALTH CHECK
# ============================================================================

@router.get("/health")
async def health_check():
    """Health check endpoint for the API"""
    return {
        "status": "healthy",
        "service": "CV Processing Pipeline",
        "version": "2.0",
        "endpoints": {
            "extract": "/api/extract",
            "enhance": "/api/enhance",
            "build": "/api/build",
            "export": "/api/export",
            "feedback": "/api/feedback"
        }
    }


@router.get("/pdf/capabilities")
async def get_pdf_capabilities():
    """
    Check PDF generation capabilities.
    
    Returns information about available PDF generation methods.
    """
    from app.services.pdf_generator import check_pdf_capabilities
    
    capabilities = check_pdf_capabilities()
    
    return {
        "success": True,
        "pdf_generation": capabilities,
        "message": "PDF generation is available" if capabilities["available"] else "No PDF backend available",
        "installation_help": {
            "pdfkit": {
                "pip": "pip install pdfkit",
                "system": "apt-get install wkhtmltopdf  # Ubuntu/Debian\nbrew install wkhtmltopdf  # macOS"
            },
            "weasyprint": {
                "pip": "pip install weasyprint",
                "system": "No system dependencies required (pure Python)"
            }
        }
    }
