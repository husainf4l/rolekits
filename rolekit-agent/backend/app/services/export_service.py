from io import BytesIO

from docx import Document
from jinja2 import Template
from weasyprint import HTML

from app.models.resume import Resume


class ExportService:
    @staticmethod
    def generate_pdf(resume: Resume, template_html: str) -> bytes:
        context = {"title": resume.title, **(resume.content or {})}
        rendered = Template(template_html).render(**context)
        html = HTML(string=rendered)
        return html.write_pdf()

    @staticmethod
    def export_docx(resume: Resume) -> bytes:
        doc = Document()

        contact = resume.content.get("contact", {})
        doc.add_heading(contact.get("full_name", resume.title), level=0)

        summary = resume.content.get("summary", {})
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(summary.get("summary", ""))

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

